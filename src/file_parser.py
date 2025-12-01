# src/file_parser.py
import pdfplumber
import docx
import os
from typing import Optional

def parse_pdf(path: str) -> str:
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text.append(t)
    return "\n".join(text)

def parse_docx(path: str) -> str:
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paragraphs)

def parse_txt(path: str) -> str:
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

def extract_text_from_file(path: str) -> Optional[str]:
    if not os.path.exists(path):
        return None
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        return parse_pdf(path)
    elif ext in ('.docx', '.doc'):
        return parse_docx(path)
    elif ext in ('.txt',):
        return parse_txt(path)
    else:
        # unsupported
        return None
def extract_text_from_bytes(data, filename):
    ext = filename.lower()

    if ext.endswith(".pdf"):
        import pdfplumber
        from io import BytesIO
        text = ""
        with pdfplumber.open(BytesIO(data)) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text

    elif ext.endswith(".docx"):
        from docx import Document
        from io import BytesIO
        doc = Document(BytesIO(data))
        return "\n".join([p.text for p in doc.paragraphs])

    else:
        return data.decode("utf-8", errors="ignore")
